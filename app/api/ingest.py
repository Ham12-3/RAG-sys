import time
import uuid

import structlog
from fastapi import APIRouter, Depends, File, Form, UploadFile

from app.config import ChunkingStrategy
from app.core.dependencies import (
    get_chunking_service,
    get_embedding_service,
    get_vectorstore_service,
)
from app.core.exceptions import DocumentParsingError
from app.models.document import Document, EmbeddedChunk
from app.schemas.ingest import ChunkInfo, IngestResponse
from app.services.chunking import ChunkingService
from app.services.embedding import EmbeddingService
from app.services.vectorstore import VectorStoreService

logger = structlog.get_logger()
router = APIRouter()


async def parse_file(file: UploadFile) -> str:
    content = await file.read()
    filename = file.filename or "unknown"

    if filename.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            import io

            reader = PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            raise DocumentParsingError(f"Failed to parse PDF: {e}") from e

    elif filename.endswith(".docx"):
        try:
            from docx import Document as DocxDocument
            import io

            doc = DocxDocument(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            raise DocumentParsingError(f"Failed to parse DOCX: {e}") from e

    else:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError as e:
            raise DocumentParsingError(f"Failed to decode file as UTF-8: {e}") from e


@router.post("/ingest", response_model=IngestResponse)
async def ingest_document(
    file: UploadFile = File(...),
    chunking_strategy: ChunkingStrategy = Form(default=ChunkingStrategy.RECURSIVE_CHARACTER),
    chunk_size: int = Form(default=512, ge=50, le=4096),
    chunk_overlap: int = Form(default=50, ge=0, le=500),
    collection_name: str | None = Form(default=None),
    chunking_service: ChunkingService = Depends(get_chunking_service),
    embedding_service: EmbeddingService = Depends(get_embedding_service),
    vectorstore: VectorStoreService = Depends(get_vectorstore_service),
):
    start = time.perf_counter()
    document_id = str(uuid.uuid4())
    source = file.filename or "upload"

    # Parse document
    text = await parse_file(file)
    if not text.strip():
        raise DocumentParsingError("Document is empty or contains no extractable text")

    document = Document(content=text, source=source)

    # Chunk
    chunks = await chunking_service.chunk(
        document, strategy=chunking_strategy, chunk_size=chunk_size, chunk_overlap=chunk_overlap
    )

    # Embed
    texts = [c.text for c in chunks]
    embeddings = await embedding_service.embed_texts(texts)

    embedded_chunks = [
        EmbeddedChunk(
            text=c.text,
            source=c.source,
            chunk_index=c.chunk_index,
            embedding=emb,
            metadata=c.metadata,
        )
        for c, emb in zip(chunks, embeddings)
    ]

    # Store
    coll_name = collection_name or vectorstore.settings.weaviate_collection_name
    await vectorstore.ensure_collection(coll_name, embedding_service.dimension)
    await vectorstore.upsert_chunks(embedded_chunks, coll_name, document_id)

    elapsed_ms = (time.perf_counter() - start) * 1000

    logger.info(
        "ingest_complete",
        document_id=document_id,
        source=source,
        chunks=len(chunks),
        elapsed_ms=round(elapsed_ms, 2),
    )

    return IngestResponse(
        document_id=document_id,
        source=source,
        total_chunks=len(chunks),
        chunks=[
            ChunkInfo(
                chunk_index=c.chunk_index,
                text_preview=c.text[:200],
                char_count=len(c.text),
            )
            for c in chunks
        ],
        embedding_provider=embedding_service.provider_name,
        processing_time_ms=round(elapsed_ms, 2),
    )
